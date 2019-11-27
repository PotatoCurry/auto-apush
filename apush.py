import docx
from gooey import Gooey, GooeyParser
from summa.summarizer import summarize

from scraper import Scraper


@Gooey(program_name="auto-apush",
       program_description="APUSH outline generator",
       progress_regex=r"^(?P<current>\d+)/(?P<total>\d+)$",
       progress_expr="current / total * 100",
       hide_progress_msg=True)
def main():
    parser = GooeyParser(description="APUSH outline generator")
    parser.add_argument("Username")
    parser.add_argument("Password", widget='PasswordField')
    parser.add_argument("Outline", widget='FileChooser')
    parser.add_argument("Summarization", widget='Dropdown', choices=['Full', '75%', '50%', '25%'])
    args = parser.parse_args()

    summarization_factors = {'Full': 1.0, '75%': 0.75, '50%': 0.5, '25%': 0.25}
    summarization_factor = summarization_factors[args.Summarization]
    outline_doc = docx.Document(args.Outline)
    print("Opened", args.Outline, flush=True)
    scraper = Scraper(args.Username, args.Password)

    section_number = 1
    output_doc = docx.Document()

    # Outline content begins at paragraph 13
    sections = outline_doc.paragraphs[13:-1]
    total_sections = sum(section.runs[0].bold is None for section in sections)

    for paragraph in outline_doc.paragraphs[13:-1]:
        section = paragraph.text.strip()
        if paragraph.runs[0].bold:
            output_doc.add_heading(section, 2)
        else:
            content = scraper.section_body(section)
            if content is None:
                print("Unable to extract body of", section, flush=True)
            else:
                content = summarize(content, summarization_factor, split=True)
                output_doc.add_paragraph(str(section_number) + ". " + section + ": " + ' '.join(content))
            print("Processed", section, flush=True)
            print(str(section_number) + '/' + str(total_sections), flush=True)
            section_number += 1

    scraper.close()
    output_filename = args.Outline[:-5] + "_Complete.docx"
    output_doc.save(output_filename)
    print("Saved completed outline as", output_filename, flush=True)


if __name__ == '__main__':
    main()
